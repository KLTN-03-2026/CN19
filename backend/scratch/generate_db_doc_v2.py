import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def parse_dbml(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    sections = []
    current_section = "General"
    tables = []
    current_table = None
    relation_map = {}

    # Concise and Easy to Understand Descriptions
    common_desc = {
        'id': 'Mã định danh duy nhất (Khóa chính).',
        'email': 'Địa chỉ email đăng nhập hệ thống.',
        'phone_number': 'Số điện thoại liên lạc.',
        'password_hash': 'Mật khẩu đã mã hóa bảo mật.',
        'role': 'Vai trò/Quyền hạn người dùng.',
        'status': 'Trạng thái hoạt động hiện tại.',
        'wallet_address': 'Địa chỉ ví Blockchain cá nhân.',
        'avatar_url': 'Ảnh đại diện người dùng.',
        'created_at': 'Thời điểm tạo bản ghi.',
        'updated_at': 'Thời điểm cập nhật bản ghi.',
        'full_name': 'Họ và tên đầy đủ.',
        'balance': 'Số dư tài khoản hệ thống.',
        'amount': 'Số tiền thực hiện giao dịch.',
        'price': 'Đơn giá niêm yết.',
        'description': 'Mô tả thông tin chi tiết.',
        'image_url': 'Đường dẫn hình ảnh minh họa.',
        'is_active': 'Trạng thái hiển thị (Bật/Tắt).',
        'title': 'Tiêu đề hiển thị chính.',
        'event_date': 'Ngày tổ chức sự kiện.',
        'stock': 'Số lượng hàng tồn trong kho.',
        'quantity': 'Số lượng mục đặt mua.',
        'unit_price': 'Giá bán thực tế mỗi đơn vị.',
        'subtotal': 'Thành tiền (chưa trừ phí).',
        'total_amount': 'Tổng tiền thanh toán cuối cùng.',
        'order_number': 'Mã số đơn hàng hệ thống.',
        'transaction_hash': 'Mã giao dịch trên Blockchain.',
        'nft_token_id': 'Mã định danh NFT duy nhất.',
        'is_bot_detected': 'Cảnh báo hành vi Bot/Gian lận.',
        'risk_score': 'Điểm rủi ro (AI đánh giá).',
        'pickup_code': 'Mã nhận hàng tại quầy.',
        'is_used': 'Đánh dấu vé đã sử dụng.',
        'payout_amount': 'Số tiền thực nhận sau phí.',
        'fee_amount': 'Số tiền phí dịch vụ.',
        'account_number': 'Số tài khoản ngân hàng.',
        'bank_name': 'Tên ngân hàng thụ hưởng.',
        'royalty_fee_percent': 'Tỷ lệ phí tác quyền (%).',
        'platform_fee': 'Phí thu bởi nền tảng.',
        'organizer_revenue': 'Doanh thu ròng nhà tổ chức.',
        'is_settled': 'Xác nhận đã quyết toán tài chính.',
        'checked_in_at': 'Thời điểm khách quét vé vào cổng.',
        'smart_contract_address': 'Địa chỉ Hợp đồng thông minh.',
        'allow_resale': 'Cho phép bán lại vé (Có/Không).',
        'is_featured': 'Đánh dấu sự kiện nổi bật.',
        'ticket_number': 'Số seri vé hệ thống.',
        'kyc_status': 'Trạng thái xác thực danh tính.',
        'organization_name': 'Tên tổ chức thực hiện.',
        'identity_card': 'Ảnh giấy tờ tùy thân.',
        'is_verified': 'Xác nhận tài khoản uy tín.'
    }

    # First pass for relations
    for line in lines:
        line = line.strip()
        if line.startswith('Ref:'):
            # Ref: Table1.col > Table2.col
            m = re.search(r'Ref:\s+(\w+)\.(\w+)\s+[\-><]\s+(\w+)\.(\w+)', line)
            if m:
                relation_map[f"{m.group(1)}.{m.group(2)}"] = f"{m.group(3)}.{m.group(4)}"

    # Second pass for structure
    for line in lines:
        line = line.strip()
        if not line: continue
        
        if line.startswith('// ---'):
            current_section = line.replace('// ---', '').replace('---', '').strip()
            sections.append({'title': current_section, 'tables': []})
            continue
            
        if line.startswith('Table'):
            m = re.match(r'Table\s+(\w+)\s+\{', line)
            if m:
                current_table = {'name': m.group(1), 'fields': []}
                if not sections: sections.append({'title': 'General', 'tables': []})
                sections[-1]['tables'].append(current_table)
            continue
            
        if line == '}':
            current_table = None
            continue
            
        if current_table and not line.startswith('Ref:'):
            # Parse field: name type [constraints]
            parts = re.split(r'\s+', line)
            if len(parts) >= 2:
                f_name = parts[0]
                f_type = parts[1]
                constraints = " ".join(parts[2:]).replace('[', '').replace(']', '')
                
                # Intelligent Description
                desc = common_desc.get(f_name, "")
                if not desc:
                    if '_id' in f_name:
                        desc = f"Khóa ngoại dùng để thiết lập mối quan hệ với bảng {f_name.replace('_id', '').capitalize()}."
                    elif 'url' in f_name: 
                        desc = "Đường dẫn lưu trữ tập tin đa phương tiện trên máy chủ đám mây."
                    # Refined time-based descriptions
                    elif f_name == 'created_at': desc = "Thời điểm ghi nhận bản ghi lần đầu vào cơ sở dữ liệu."
                    elif f_name == 'updated_at': desc = "Thời điểm cập nhật các thay đổi mới nhất của bản ghi."
                    elif f_name == 'paid_at': desc = "Thời điểm hệ thống xác nhận thanh toán thành công từ cổng thanh toán."
                    elif f_name == 'sold_at': desc = "Thời điểm hoàn tất chuyển giao sở hữu tài sản cho người mua mới."
                    elif f_name == 'checked_in_at': desc = "Thời điểm khách hàng quét mã QR để vào cổng sự kiện."
                    elif f_name == 'kyc_verified_at': desc = "Thời điểm bộ phận quản trị phê duyệt hồ sơ pháp lý của tổ chức."
                    elif f_name == 'expires_at': desc = "Thời điểm hết hạn hiệu lực của đơn hàng hoặc phiên làm việc."
                    elif f_name == 'requested_at': desc = "Thời điểm người dùng gửi yêu cầu (hoàn tiền, chuyển vé) lên hệ thống."
                    elif f_name == 'completed_at': desc = "Thời điểm quy trình nghiệp vụ được xử lý hoàn tất 100%."
                    elif f_name == 'processed_at': desc = "Thời điểm máy chủ hoàn tất các tác vụ tính toán và lưu trữ liên quan."
                    elif f_name == 'redeemed_at': desc = "Thời điểm vật phẩm quà tặng được trao cho người sở hữu hợp lệ."
                    elif f_name == 'scanned_at': desc = "Thời điểm thực hiện thao tác kiểm tra tính hợp lệ của vé qua ứng dụng quét mã."
                    elif 'at' in f_name.lower(): desc = f"Mốc thời gian ghi nhận trạng thái {f_name.replace('_at', '')} trong quy trình."
                    else: desc = f"Lưu trữ thông tin bổ trợ cho thuộc tính {f_name.replace('_', ' ')} của thực thể."
                
                current_table['fields'].append({
                    'Column': f_name,
                    'Type': f_type.upper(),
                    'Null': "No" if 'pk' in constraints or 'not null' in constraints.lower() else "Yes",
                    'Extra': constraints,
                    'Link': relation_map.get(f"{current_table['name']}.{f_name}", ""),
                    'Desc': desc
                })

    return sections

def generate_doc(sections, output_path):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Inches

    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    title = doc.add_heading('HỆ THỐNG QUẢN LÝ VÉ BASTICKET - PHÂN TÍCH CƠ SỞ DỮ LIỆU', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16) # Title slightly larger

    table_counter = 1
    chapter_number = 3 

    type_map_v2 = {
        'UUID': 'uuid',
        'VARCHAR': 'varchar(255)',
        'INT': 'int',
        'DATETIME': 'datetime',
        'DECIMAL': 'decimal(18,2)',
        'TEXT': 'text',
        'BOOLEAN': 'bit(1)',
        'FLOAT': 'float',
        'JSON': 'jsonb',
    }

    # Define Fixed Widths (Total ~6.5 inches for A4)
    col_widths = [
        Inches(1.0), # Column
        Inches(1.1), # Type
        Inches(0.5), # Null
        Inches(0.8), # Extra
        Inches(0.9), # Link to
        Inches(2.2), # Mô tả (Wide)
    ]

    for section in sections:
        if not section['tables']: continue
        
        h = doc.add_heading(section['title'], level=1)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        for table in section['tables']:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(f'Bảng {chapter_number}.{table_counter} Bảng {table["name"]}')
            run_title.font.name = 'Times New Roman'
            run_title.italic = True
            run_title.font.size = Pt(13)
            table_counter += 1
            
            t = doc.add_table(rows=1, cols=6)
            t.style = 'Table Grid'
            t.autofit = False
            
            # Header
            hdr_cells = t.rows[0].cells
            headers = ['Column', 'Type', 'Null', 'Extra', 'Link to', 'Mô tả']
            for i, h_text in enumerate(headers):
                hdr_cells[i].text = h_text
                # Force width on header cell
                hdr_cells[i].width = col_widths[i]
                hdr_cells[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w'))))
                p = hdr_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.bold = True
                run.font.size = Pt(12)
            
            # Data
            for f in table['fields']:
                row = t.add_row()
                row_cells = row.cells
                row_cells[0].text = f['Column']
                
                raw_t = f['Type'].replace('"', '')
                row_cells[1].text = type_map_v2.get(raw_t, raw_t)
                
                row_cells[2].text = f['Null']
                row_cells[3].text = f['Extra']
                row_cells[4].text = f['Link']
                row_cells[5].text = f['Desc']
                
                for i, cell in enumerate(row_cells):
                    # Force width on EVERY data cell
                    cell.width = col_widths[i]
                    for paragraph in cell.paragraphs:
                        # Ensure cell margins/padding are consistent
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                
                # Make Column name bold
                p_col = row_cells[0].paragraphs[0]
                if p_col.runs:
                    p_col.runs[0].bold = True
            
            doc.add_paragraph() 

    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    input_f = r'd:\KLTN_HeThongVeSuKien\BASTICKET\backend\scratch\db_schema_input.dbml'
    output_f = r'd:\KLTN_HeThongVeSuKien\BASTICKET\backend\docs\BASTICKET_Database_Perfect_Report.docx'
    
    data = parse_dbml(input_f)
    generate_doc(data, output_f)
